"""添付ファイルの取り込み（テキスト抽出）と書き出し（返信への添付）。"""

from __future__ import annotations

import csv
import io
import json
import logging
import mimetypes
import re
from dataclasses import dataclass

log = logging.getLogger(__name__)

# 拡張子だけでテキストとみなすもの（MIME が application/octet-stream で届くことが多いため）
TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".tsv", ".json", ".jsonl",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".env", ".sql", ".html", ".htm",
    ".xml", ".svg", ".py", ".js", ".mjs", ".ts", ".tsx", ".jsx", ".c", ".h", ".cpp",
    ".hpp", ".cs", ".java", ".kt", ".go", ".rs", ".rb", ".php", ".swift", ".sh", ".bash",
    ".zsh", ".ps1", ".r", ".m", ".jl", ".lua", ".pl", ".vue", ".css", ".scss", ".diff",
    ".patch", ".gitignore", ".dockerfile",
}

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff", ".heic"}

# 抽出したテキストが長すぎる添付を 1 件あたりで切り詰める上限
PER_ATTACHMENT_CHAR_LIMIT = 40_000


@dataclass
class ExtractedAttachment:
    """受信した添付 1 件と、そこから取り出したテキスト。"""

    filename: str
    content_type: str
    size: int
    text: str | None
    note: str | None = None

    def as_prompt_block(self, char_limit: int = PER_ATTACHMENT_CHAR_LIMIT) -> str:
        header = f"### 添付ファイル: {self.filename} ({self.content_type}, {self.size:,} バイト)"
        if self.text is None:
            return f"{header}\n{self.note or '（本文を抽出できませんでした）'}"
        body = self.text
        truncated = ""
        if len(body) > char_limit:
            body = body[:char_limit]
            truncated = f"\n…（以降 {self.size:,} バイト中の残りは省略）"
        note = f"\n{self.note}" if self.note else ""
        return f"{header}{note}\n```\n{body}{truncated}\n```"


@dataclass
class OutgoingFile:
    """返信に添付するファイル。"""

    filename: str
    content: bytes

    @property
    def content_type(self) -> str:
        guessed, _ = mimetypes.guess_type(self.filename)
        return guessed or "application/octet-stream"


def _suffix(filename: str) -> str:
    _, _, ext = filename.lower().rpartition(".")
    return f".{ext}" if ext and ext != filename.lower() else ""


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp932", "euc-jp", "iso-2022-jp", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> tuple[str | None, str | None]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return None, "（PDF を読むには pypdf をインストールしてください）"
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"--- {index} ページ ---\n{text}")
        if not pages:
            return None, "（テキストを含まない PDF のようです。スキャン画像の可能性があります）"
        return "\n\n".join(pages), None
    except Exception as exc:  # noqa: BLE001 - 壊れた PDF でボットを止めない
        log.warning("PDF の解析に失敗しました: %s", exc)
        return None, f"（PDF の解析に失敗しました: {exc}）"


def _extract_docx(data: bytes) -> tuple[str | None, str | None]:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError:
        return None, "（Word 文書を読むには python-docx をインストールしてください）"
    try:
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts) or None, None if parts else "（本文が空でした）"
    except Exception as exc:  # noqa: BLE001
        log.warning("Word 文書の解析に失敗しました: %s", exc)
        return None, f"（Word 文書の解析に失敗しました: {exc}）"


def _extract_xlsx(data: bytes) -> tuple[str | None, str | None]:
    try:
        import openpyxl
    except ImportError:
        return None, "（Excel を読むには openpyxl をインストールしてください）"
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001
        log.warning("Excel の解析に失敗しました: %s", exc)
        return None, f"（Excel の解析に失敗しました: {exc}）"

    sheets: list[str] = []
    max_rows = 500
    try:
        for sheet in workbook.worksheets:
            buffer = io.StringIO()
            writer = csv.writer(buffer)
            rows = 0
            for row in sheet.iter_rows(values_only=True):
                if rows >= max_rows:
                    buffer.write(f"...（{sheet.title} は先頭 {max_rows} 行のみ）\n")
                    break
                if row is None or all(cell is None for cell in row):
                    continue
                writer.writerow(["" if cell is None else str(cell) for cell in row])
                rows += 1
            if rows:
                sheets.append(f"--- シート: {sheet.title} ---\n{buffer.getvalue()}")
    finally:
        workbook.close()

    if not sheets:
        return None, "（データのあるシートが見つかりませんでした）"
    return "\n".join(sheets), "（CSV 形式に変換して渡しています）"


def extract_attachment(filename: str, content_type: str, data: bytes) -> ExtractedAttachment:
    """添付 1 件を DeepSeek に渡せるテキストへ変換する。"""

    filename = filename or "attachment"
    suffix = _suffix(filename)
    lowered = content_type.lower()
    base = ExtractedAttachment(filename=filename, content_type=content_type, size=len(data), text=None)

    if suffix == ".pdf" or "pdf" in lowered:
        base.text, base.note = _extract_pdf(data)
        return base

    if suffix in {".docx", ".dotx"} or "wordprocessingml" in lowered:
        base.text, base.note = _extract_docx(data)
        return base

    if suffix in {".xlsx", ".xlsm", ".xltx"} or "spreadsheetml" in lowered:
        base.text, base.note = _extract_xlsx(data)
        return base

    if suffix in IMAGE_SUFFIXES or lowered.startswith("image/"):
        base.note = "（DeepSeek は画像を読めないため、この添付は本文に含めていません）"
        return base

    if suffix in TEXT_SUFFIXES or lowered.startswith("text/") or lowered in {
        "application/json",
        "application/xml",
        "application/javascript",
        "application/x-yaml",
    }:
        base.text = _decode_text(data)
        return base

    if suffix in {".doc", ".xls", ".ppt", ".pptx"}:
        base.note = f"（{suffix} 形式には未対応です。PDF か docx/xlsx に変換して再送してください）"
        return base

    # 未知の形式でも、UTF-8 として素直に読めるならテキスト扱いにする。
    try:
        decoded = data.decode("utf-8")
    except UnicodeDecodeError:
        base.note = "（バイナリのため本文を抽出できませんでした）"
        return base
    if "\x00" in decoded:
        base.note = "（バイナリのため本文を抽出できませんでした）"
        return base
    base.text = decoded
    return base


_FENCE_RE = re.compile(r"^(?P<indent>[ \t]*)(?P<fence>`{3,}|~{3,})[ \t]*(?P<info>[^\n`]*)$")
_FILE_INFO_RE = re.compile(r"^(?:file|attachment|添付)\s*[:：]\s*(?P<name>.+?)\s*$", re.IGNORECASE)

_UNSAFE_NAME_RE = re.compile(r"[^A-Za-z0-9._\-぀-ヿ㐀-鿿]+")


def sanitize_filename(name: str, fallback: str = "attachment.txt") -> str:
    """パス区切りや制御文字を落として、安全なファイル名にする。"""

    name = name.strip().replace("\\", "/").split("/")[-1]
    name = name.strip().strip(".")
    name = _UNSAFE_NAME_RE.sub("_", name).strip("_")
    if not name:
        return fallback
    return name[:120]


def split_file_blocks(text: str) -> tuple[str, list[OutgoingFile]]:
    """本文から ```file:NAME フェンスを抜き出し、(残った本文, 添付リスト) を返す。"""

    lines = text.splitlines()
    body_lines: list[str] = []
    files: list[OutgoingFile] = []
    used_names: dict[str, int] = {}

    index = 0
    while index < len(lines):
        match = _FENCE_RE.match(lines[index])
        info_match = _FILE_INFO_RE.match(match.group("info").strip()) if match else None
        if not match or not info_match:
            body_lines.append(lines[index])
            index += 1
            continue

        fence = match.group("fence")
        indent = match.group("indent")
        content: list[str] = []
        index += 1
        closed = False
        while index < len(lines):
            candidate = lines[index].strip()
            if candidate.startswith(fence[0] * len(fence)) and set(candidate) == {fence[0]}:
                closed = True
                index += 1
                break
            line = lines[index]
            if indent and line.startswith(indent):
                line = line[len(indent):]
            content.append(line)
            index += 1

        if not closed:
            log.warning("閉じられていない file フェンスを検出したため本文として扱います。")
            body_lines.append(match.group(0))
            body_lines.extend(content)
            continue

        raw_name = info_match.group("name")
        filename = sanitize_filename(raw_name)
        if filename in used_names:
            used_names[filename] += 1
            stem, dot, ext = filename.rpartition(".")
            suffix = f"-{used_names[filename]}"
            filename = f"{stem}{suffix}{dot}{ext}" if dot else f"{filename}{suffix}"
        else:
            used_names[filename] = 1

        payload = "\n".join(content)
        if payload and not payload.endswith("\n"):
            payload += "\n"
        files.append(OutgoingFile(filename=filename, content=payload.encode("utf-8")))
        body_lines.append(f"［添付: {filename}］")

    body = "\n".join(body_lines).strip()
    return body, files


def summarize_json(value: object) -> str:
    """ログ用に構造化データを短く文字列化する。"""

    try:
        return json.dumps(value, ensure_ascii=False)[:500]
    except (TypeError, ValueError):
        return repr(value)[:500]
